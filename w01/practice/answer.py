#!/usr/bin/env python3

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Q1. inputフォルダの中にある'.txt'のファイルを指定しよう
guest_file = open('./input/guests.txt', 'r', encoding='utf-8')

doc = docx.Document()

for guest in guest_file:
    guest = guest.strip()
    if guest == '':
        continue
    doc.add_paragraph(guest)
    doc.add_paragraph('拝啓　時下ますますご盛栄のこととお喜び申し上げます。')
    doc.add_paragraph('このたび下記の通りパーティーを開催しますので、' +
                      'ご出席賜りますようよろしくお願い申し上げます。')
    # Q2. '敬具'を右寄せに指定しよう
    doc.add_paragraph('敬具').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Q3. '記'を中央寄せに指定しよう
    doc.add_paragraph('記').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\t日時：4月1日　19:00')
    doc.add_paragraph('\t会場：秋田県立大学')
    # Q4. '以上'を右寄せに指定しよう
    doc.add_paragraph('以上').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

os.mkdir('output')
doc.save('./output/invitations.docx')
print('招待状をoutputフォルダに保存しました')
